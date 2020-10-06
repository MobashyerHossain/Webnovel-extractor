from bs4 import BeautifulSoup as soup
from docx import Document
import regex as re
import requests
import os

########################################     Helpers Start     ########################################

def options(limit):
    option = -1
    while(option < 1 or option > limit):
        option = int(input('\nchoice: '))
    
    return option

def clear():
    os.system('cls')

def get_page_html(link):
    page = requests.get(link)
    page_html = soup(page.content, "html.parser")
    page.close()
    return page_html

def check_novel_in_favourite(novel_given):
    status = False
    favs = []
    try:
        f = open("wuxia_favourites_list.txt", 'r')
        favs = f.readlines()
        f.close()
    except:
        pass
    
    if (len(favs) != 0):
        if novel_given in favs:
            status = True

    return status

def structuring_chapter(cpt):
    #replace all unneccessary content
    cpt = cpt.replace("<br/>", "").replace("Translator: MrVoltaire1  ","").replace("Editor: Modlawls123","")
    cpt = cpt.replace("ChapterMid();","").replace("<script>","").replace("</script>","")
    cpt = cpt.replace('<div class="Readarea ReadAjax_content" id="chaptercontent">', "")
    cpt = cpt.replace('<div style="width:340;height:55px;overflow:hidden">', "")
    cpt = cpt.replace('<ins class="adsbygoogle" data-ad-client="ca-pub-2853920792116568" data-ad-slot="8807674870" style="display:inline-block;width:336px;height:50px">', "")
    cpt = cpt.replace('</ins>', "")
    cpt = cpt.replace('</div>', "")
    cpt = cpt.replace('Find authorized novels in Webnovel，faster updates, better experience，Please click www.webnovel.com  for visiting.', "")
    cpt = cpt.replace("(adsbygoogle = window.adsbygoogle || []).push({});", "")
    cpt = cpt.replace("Previous Chapter", "").replace("Next Chapter", "")    
    cpt = cpt.replace("Find authorized novels in Webnovel, faster updates, better experience, Please click for visiting.", "")

    # for Omniscient Reader
    cpt = cpt.replace("-ssi", "")
    cpt = cpt.replace("-nim", "")
    cpt = cpt.replace("Lv.", "level")
    cpt = cpt.replace("lv.", "level")
    cpt = cpt.replace("Ho-Seong", "Hosung")
    cpt = cpt.replace("Ho-seong", "Hosung")
    cpt = cpt.replace("Ha-Yeong", "Hayoung")
    cpt = cpt.replace("Ha-yeong", "Hayoung")
    cpt = cpt.replace("Jeong Hui-Won", "Jung Heewon")
    cpt = cpt.replace("Hui-Won", "Heewon")
    cpt = cpt.replace("Hui-won", "Heewon")
    cpt = cpt.replace("Yi Hyeon-Seong", "Lee Hyunsung")
    cpt = cpt.replace("Hyeon-Seong", "Hyunsung")
    cpt = cpt.replace("Hyeon-seong", "Hyunsung")
    cpt = cpt.replace("Yu Joong-Hyeok", "Yu Jungyeok")
    cpt = cpt.replace("Joong-Hyeok", "Jonghyuk")
    cpt = cpt.replace("Joong-hyeok", "Jonghyuk")
    cpt = cpt.replace("Jung-Hyeok", "Jonghyuk")
    cpt = cpt.replace("Jung-hyeok", "Jonghyuk")
    cpt = cpt.replace("Yi Gil-Yeong", "Lee Gilyoung")
    cpt = cpt.replace("Gil-Yeong", "Gilyoung")
    cpt = cpt.replace("Gil-yeong", "Gilyoung")
    cpt = cpt.replace("Shin Yu-Seung", "Shin Yoosung")
    cpt = cpt.replace("Yu-Seung", "Yoosung")
    cpt = cpt.replace("Yu-seung", "Yoosung")
    cpt = cpt.replace("Mister Dok-Ja", "Kim Dokja")
    cpt = cpt.replace("Dok-Ja", "Dokja")
    cpt = cpt.replace("Dok-ja", "Dokja")
    cpt = cpt.replace("Yu Sang-Ah", "Yoo Sangah")
    cpt = cpt.replace("Sang-Ah", "Sangah")
    cpt = cpt.replace("Sang-ah", "Sangah")
    cpt = cpt.replace("Han Myeong-Oh", "Han Myungoh")
    cpt = cpt.replace("Han Su-Yeong", "Han Suyeong")
    cpt = cpt.replace("Su-Yeong", "Suyeong")
    cpt = cpt.replace("Su-yeong", "Suyeong")
    cpt = cpt.replace("Myeong-Oh", "Myungoh")
    cpt = cpt.replace("Myeong-oh", "Myungoh")
    cpt = cpt.replace("Lightning Transformation", "Electrification")
    cpt = cpt.replace("Lamarck's Giraffe", "Lamarck's Kirin")
    cpt = cpt.replace("Yi Ji-Hye", "Lee Jihye")
    cpt = cpt.replace("Ji-Hye", "Jihye")
    cpt = cpt.replace("Ji-hye", "Jihye")
    cpt = cpt.replace("mantra", "True voice")
    cpt = cpt.replace("Lily Blooming on Aquarius", "Lily Pin of Aquarius")
    cpt = cpt.replace("cross-legged", "cross legged")
    cpt = cpt.replace("much-younger", "much younger")
    cpt = cpt.replace("Cheok Jun-Gyeong", "Cheok Jungyeong")
    cpt = cpt.replace("Jun-Gyeong", "Jungyeong")
    cpt = cpt.replace("Jun-gyeong", "Jungyeong")
    cpt = cpt.replace("Jeon Woo-Chi", "Jeon Woochi")
    cpt = cpt.replace("Woo-Chi", "Woochi")
    cpt = cpt.replace("Yi Seol-Hwa", "Lee Seolhwa")
    cpt = cpt.replace("Seol-Hwa", "Seolhwa")
    cpt = cpt.replace("Seol-hwa", "Seolhwa")
    cpt = cpt.replace("Yu-Shin", "Yushin")
    cpt = cpt.replace("Yu-shin", "Yushin")
    cpt = cpt.replace("Hak-Hyeon", "Hakhyeon")
    cpt = cpt.replace("Hak-hyeon", "Hakhyeon")
    cpt = cpt.replace("&lt;kim dok-ja=\"\" company=\"\"&gt;'s", "Kim Dokja's Company")
    cpt = cpt.replace("dok-ja", "dokja")    
    cpt = cpt.replace("Fable", "Story")

    #transform multiple emptylines to single empty line
    cpt = re.sub(r'\n\s*\n ', '\n\n', cpt)
    cpt = re.sub(r'\<[^<>]*\>', '', cpt)


    #uncensoring censored words
    cpt = cpt.replace("f*ck", "fuck").replace("sh*t", "shit").replace("*ss", "ass").replace("b*stard", "bastard")
    cpt = cpt.replace("b*tch", "bitch").replace("d*mn", "damn").replace("fu*k", "fuck")
    cpt = cpt.replace("F*ck", "Fuck").replace("Sh*t", "Shit").replace("*ss", "ass").replace("B*stard", "Bastard")
    cpt = cpt.replace("B*tch", "Bitch").replace("D*mn", "Damn").replace("Fu*k", "Fuck")

    return cpt

#########################################     Helpers End     #########################################

########################################   Wuxia World Start   ########################################

def wuxia_world():
    clear()

    site_link = "https://m.wuxiaworld.co"

    print('Welcome to Wuxia World!\n')
    print('1. Choose from Faveourites')
    print('2. Search by Novel name')
    print('3. Go Back')

    option = options(3)

    if option == 1:
        wuxia_faveourites(site_link)
    elif option == 2:
        wuxia_search(site_link)
    else:
        main()

def wuxia_search(site_link, novel_name = ''):
    clear()
    print("Search Page!\n")

    if novel_name == '':
        novel_name = input("Novel Name: ")
    
    novel = novel_name.lower().replace(' ', '-').replace("'", '-').replace('\n', '')
    novel_page_link = '/' +  novel + '/'

    novel_page_html = get_page_html(site_link + novel_page_link)
    chapter_list = []

    try:
        chapter_list = novel_page_html.find_all("a", class_="chapter-item")
    except:
        pass

    total_chapters = len(chapter_list)

    if(total_chapters == 0):
        print("Novel Not Found!")
        print("\n1. Try Another Name")
        print("2. Go Back")
        option = options(2)
        
        if option == 1:
            wuxia_search(site_link)
        else:
            wuxia_world()        
    else:
        print("Novel '{}' Found!".format(novel_name))
        if check_novel_in_favourite(novel_name):
            print("Novel already Added to Favourites")

        print("\n1. View Novel Detail")

        if not check_novel_in_favourite(novel_name):
            print("2. Add Novel to Favourites")
            print("3. Go Back")
        else:
            print("2. Go Back")             
        
        if not check_novel_in_favourite(novel_name):
            option = options(3)

            if option == 1:
                wuxia_novel_detail(novel_name, site_link, novel_page_link, 1)
            elif option == 2:
                f = open("wuxia_favourites_list.txt", 'a')
                f.write("\n{}".format(novel_name))
                f.close()
                wuxia_novel_detail(novel_name, site_link, novel_page_link, 0)
            else:
                wuxia_world()
        else:
            option = options(2)

            if option == 1:
                wuxia_novel_detail(novel_name, site_link, novel_page_link, 0)
            else:
                wuxia_world()

def wuxia_faveourites(site_link):
    clear()
    print("Favourites!\n")

    try:
        f = open("wuxia_favourites_list.txt", 'r')
        favs = [x.replace('\n', '') for x in f.readlines() if x != '\n']
        f.close()
    except:
        favs = []    

    if (len(favs) != 0):
        for i, line in enumerate(favs, 1):
            print("{}. {}".format(i, line.replace('\n', '')))
        print('{}. Go Back'.format(len(favs)+1))
    else:
        print('No Favourites Added yet')
        print('1. Go Back')
        option = options(1)

        if option == 1:
            wuxia_world()    

    option = options(len(favs)+1)  

    if option == len(favs)+1:
        wuxia_world()
    else:
        novel_name = favs[option-1]
        novel = novel_name.lower().replace(' ', '-').replace("'", '-').replace('\n', '')
        novel_page_link = '/' +  novel + '/'
        wuxia_novel_detail(novel_name, site_link, novel_page_link)

def wuxia_novel_detail(novel_name, site_link, novel_page_link, backtrack = 0):
    clear()
    print("{}\n".format(novel_name))

    print("1. Browse Chapters")
    print("2. View Summary")
    print("3. Go Back")

    option = options(3)

    if option == 1:
        wuxia_novel(novel_name, site_link, novel_page_link, backtrack)
    else:
        if backtrack == 0:
            wuxia_faveourites(site_link)
        else:
            wuxia_search(site_link, novel_name)


def wuxia_novel(novel_name, site_link, novel_page_link, backtrack = 0):
    clear()
    print("Chapter List of\n{}\n".format(novel_name))
    # print(novel_page_link)

    novel_page_html = get_page_html(site_link + novel_page_link)
    chapter_list = novel_page_html.find_all("a", class_="chapter-item")
    total_chapters = len(chapter_list)


    for i, chapter in enumerate(chapter_list, 1):
        print(chapter.get_text())

    print("\nWhat would you like to do?")
    print("1. Save All Chapters")
    print("2. Save Chapters in Selected Range")
    print("3. Save from 'X' Chapter to the Last chapter")
    print("4. Save from First Chapter to 'X' Chapter")
    print("5. Go Back")

    option = options(5)
    
    start = [0,0]
    end = [0,0]

    if option == 1:
        # save all chapters
        start = [get_chapter_index(chapter_list, 0), 0]
        end = [get_chapter_index(chapter_list, total_chapters-1), total_chapters-1]
        
    elif option == 2:
        # save selected range of chapters
        list_start = -1
        cpt_start = -1
        while(list_start < 1 or list_start > (total_chapters - 1)):
            cpt_start = int(input('Choose a Starting Chapter: '))
            # chapter index check
            list_start = get_chapter_list_index(chapter_list, cpt_start)
        start = [cpt_start, list_start]

        list_end = -1
        cpt_end = -1
        while(list_end < list_start or list_end > (total_chapters - 1)):
            cpt_end = int(input('Choose a Ending Chapter: '))
            # chapter index check
            list_end = get_chapter_list_index(chapter_list, cpt_end)
        end = [cpt_end, list_end]

    elif option == 3:
        # save from X-Last
        list_start = -1
        cpt_start = -1
        while(list_start < 1 or list_start > (total_chapters - 1)):
            cpt_start = int(input('Choose a Starting Chapter: '))
            # chapter index check
            list_start = get_chapter_list_index(chapter_list, cpt_start)
        start = [cpt_start, list_start]
        end = [get_chapter_index(chapter_list, total_chapters-1), total_chapters-1]

    elif option == 4:
        # save from X-Last
        list_end = -1
        cpt_end = -1
        while(list_end < 1 or list_end > (total_chapters - 1)):
            cpt_end = int(input('Choose a Ending Chapter: '))
            # chapter index check
            list_end = get_chapter_list_index(chapter_list, cpt_end)
        end = [cpt_end, list_end]
        start = [get_chapter_index(chapter_list, 0), 0]
 
    else:
        if backtrack == 0:
            wuxia_faveourites(site_link)
        else:
            wuxia_search(site_link, novel_name)
    
    save_chapters(novel_name, site_link, novel_page_link, chapter_list, start, end)
        
def save_chapters(novel_name, site_link, novel_page_link, chapter_list, start, end):
    clear()
    print("Chapter are being Saved!!\n")

    # Create target Directory if don't exist
    dirName = novel_name.replace("'", '').replace('\n', '')
    if not os.path.exists(dirName):
        os.mkdir(dirName)
        print("Directory " , dirName ,  " Created ")

    #document initialization
    LNdocument = Document()

    st_cpt = start[0]

    for index, chapter_link in enumerate(chapter_list[start[1]:end[1]+1], start[0]):
        chapter_link = site_link + chapter_link['href']

        # print title
        chapter_title, chapter_content = get_chapter_content(chapter_link)
        print(chapter_title)

        # add to document
        LNdocument.add_heading(chapter_title, level=1)
        LNdocument.add_paragraph(chapter_content)

        #devide every 50 chapter in different doc file
        if (index % 50 == 0):            
            doc_name = "{} {} - {}.docx".format(dirName, str(st_cpt), str(index))
            LNdocument.save(dirName+'/'+doc_name)
            LNdocument = Document()

            st_cpt = index + 1

    doc_name = "{} {} - {}.docx".format(dirName, str(st_cpt), str(end[0]))
    LNdocument.save(dirName+'/'+doc_name)

    print('\n')
    print("1. Go Back to Chapter List")
    print("2. Go Back to Novel Detail")
    print("3. Go back to Favourites")
    print("4. Go Back to Site Selection")

    option = options(4)

    if option == 1:
        wuxia_novel(novel_name, site_link, novel_page_link)
    elif option == 2:
        wuxia_novel_detail(novel_name, site_link, novel_page_link)
    elif option == 3:
        wuxia_faveourites(site_link)
    else:
        main()

def get_chapter_content(chapter_link):
    chapter_page_html = get_page_html(chapter_link)
    # print(chapter_page_html.prettify())
    chapter_title = chapter_page_html.find("h1", class_="chapter-title").get_text()
    chapter_content = chapter_page_html.find("div", class_="chapter-entity")
    chapter_content = structuring_chapter(chapter_content.prettify())

    return chapter_title, chapter_content

def get_chapter_list_index(chapter_list, cpt_index):
    index = 0
    for i, cpt in enumerate(chapter_list, 0):
        if str(cpt_index) in cpt.get_text():
            index = i
            break
    
    return index

def get_chapter_index(chapter_list, list_index):
    chapter = chapter_list[list_index].get_text()
    chapter = chapter.replace('.', ' ').replace('-', ' ')
    index = 0
    for i in chapter.split(' '):
        try:
            index = int(i)
        except:
            continue
        break

    return int(index)


    

########################################    Wuxia World End    ########################################

########################################    Novel Full Start   ########################################

def novel_full():
    clear()
    return 1

########################################     Novel Full End    ########################################

###########################################    Main Start   ###########################################

def main():
    clear()

    print("Choose your preferred Website?")
    print("1. Wuxiaworld")
    print("2. Novelfull")

    site_choice = options(2)

    if site_choice == 1:
        wuxia_world()
    else:
        novel_full()

if __name__=="__main__":
    main()

###########################################     Main End    ###########################################